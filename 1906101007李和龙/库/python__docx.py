import docx
"""创建文档"""
document = docx.Document()
"""添加标题"""
document.add_heading("合并后",0)
"""读取要合并的文档"""
document1 = docx.Document("20190922-李和龙-关于python部分基础语法的概述.docx")
document2 = docx.Document("20190929-李和龙-用python解决阶乘问题.docx")
"""遍历每段并加起来"""
for par1 in document1.paragraphs:
    document.add_paragraph(par1.text)
"""加换页"""
document.add_page_break()
"""继续遍历"""
for par2 in document2.paragraphs:
    document.add_paragraph(par2.text)
"""字体调整"""

"""保存文档"""
document.save("文档合并.docx")