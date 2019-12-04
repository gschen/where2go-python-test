#创建一个docx文档，并编辑标题，字体，颜色，对齐格式
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import RGBColor
document = Document('jiang.docx')
p = document.add_paragraph()
run = p.add_run('八里公路')
run.font.size = Pt(26)
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
run.bold = True
document.styles['Normal'].font.name = '宋体'
run.font.color.rgb = RGBColor(255,0,0)
document.save('jiang.docx')